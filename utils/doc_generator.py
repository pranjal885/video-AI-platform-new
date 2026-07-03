from docx import Document
import os


def generate_doc(summary, filename):

    output_folder = "outputs"
    os.makedirs(output_folder, exist_ok=True)

    doc_path = os.path.join(
        output_folder,
        filename + ".docx"
    )

    document = Document()

    document.add_heading(
        "AI Generated Summary",
        level=1
    )

    document.add_paragraph(summary)

    document.save(doc_path)

    return doc_path